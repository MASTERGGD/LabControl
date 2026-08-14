from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "manuales" / "Manual_docente_expediente_academico.docx"

BLUE = "2563EB"
NAVY = "17365D"
LIGHT_BLUE = "EAF2FF"
PALE = "F4F7FB"
GRAY = "667085"
GREEN = "15803D"
AMBER = "B45309"
RED = "B42318"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc.get_or_add_tcPr()
    tc_mar = tc.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep(paragraph, next_=False):
    ppr = paragraph._p.get_or_add_pPr()
    name = "keepNext" if next_ else "keepLines"
    if ppr.find(qn(f"w:{name}")) is None:
        ppr.append(OxmlElement(f"w:{name}"))


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    keep(p, True)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    keep(p)
    return p


def add_step(doc, number, title, body):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    t.columns[0].width = Inches(0.55)
    t.columns[1].width = Inches(5.95)
    c0, c1 = t.rows[0].cells
    c0.width, c1.width = Inches(0.55), Inches(5.95)
    shade(c0, BLUE)
    c0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(str(number))
    r0.bold = True
    r0.font.color.rgb = RGBColor(255, 255, 255)
    p1 = c1.paragraphs[0]
    r = p1.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p1.add_run("\n" + body)
    for c in (c0, c1): margins(c, 110, 130, 110, 130)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def callout(doc, label, text, fill=LIGHT_BLUE, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    shade(cell, fill)
    margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    r = p.add_run(label + "  ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("263238")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    for name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, NAVY),
        ("Heading 3", 11.5, 9, 4, NAVY),
    ):
        s = styles[name]
        s.font.name = "Aptos Display"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        styles[name].font.name = "Aptos"
        styles[name].font.size = Pt(10.5)
        styles[name].paragraph_format.left_indent = Inches(0.38)
        styles[name].paragraph_format.first_line_indent = Inches(-0.19)
        styles[name].paragraph_format.space_after = Pt(4)
        styles[name].paragraph_format.line_spacing = 1.18

    # Header/footer
    hp = sec.header.paragraphs[0]
    hp.text = "SIGA  |  Manual de usuario"
    hp.style = normal
    hp.runs[0].font.size = Pt(8.5)
    hp.runs[0].font.bold = True
    hp.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    add_page_number(sec.footer.paragraphs[0])

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("MANUAL DE USUARIO")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    r.font.letter_spacing = Pt(1.2)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    rt = title.add_run("Expediente Académico Integral")
    rt.bold = True
    rt.font.name = "Aptos Display"
    rt.font.size = Pt(30)
    rt.font.color.rgb = RGBColor.from_string(NAVY)
    sub = doc.add_paragraph("Guía para docentes y tutores")
    sub.paragraph_format.space_after = Pt(24)
    sub.runs[0].font.size = Pt(16)
    sub.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    callout(doc, "PROPÓSITO", "Consultar el panorama académico de los grupos y analizar, en un solo lugar, la trayectoria, asistencia, evidencias, acuerdos y seguimiento tutorial de cada estudiante.")
    doc.add_paragraph()
    meta = doc.add_table(rows=3, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    meta.columns[0].width = Inches(1.65)
    meta.columns[1].width = Inches(4.85)
    for i, (label, value) in enumerate((("Perfil", "Docente / tutor"), ("Sistema", "SIGA - Seguimiento institucional"), ("Versión", "1.0 · Agosto de 2026"))):
        c0, c1 = meta.rows[i].cells
        c0.width, c1.width = Inches(1.65), Inches(4.85)
        shade(c0, PALE)
        c0.paragraphs[0].add_run(label).bold = True
        c1.paragraphs[0].add_run(value)
        for c in (c0, c1): margins(c)
    doc.add_paragraph()
    p = doc.add_paragraph("Uso institucional. La información del expediente es confidencial y debe emplearse exclusivamente para fines de acompañamiento académico.")
    p.runs[0].italic = True
    p.runs[0].font.color.rgb = RGBColor.from_string(GRAY)

    doc.add_page_break()
    add_heading(doc, "Contenido", 1)
    for n, item in enumerate((
        "Objetivo y alcance", "Ingreso al expediente", "Panorama del grupo", "Consulta del expediente de un alumno",
        "Interpretación de las secciones", "Flujo recomendado de seguimiento", "Buenas prácticas y confidencialidad",
        "Problemas frecuentes", "Guía rápida"
    ), 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)
    callout(doc, "LECTURA RÁPIDA", "Si es la primera vez que utiliza el módulo, revise primero “Ingreso al expediente”, “Panorama del grupo” y “Flujo recomendado de seguimiento”.")

    add_heading(doc, "1. Objetivo y alcance", 1)
    doc.add_paragraph("El Expediente Académico Integral reúne indicadores procedentes de los registros disponibles en SIGA. Su finalidad es apoyar la detección oportuna de estudiantes que requieren acompañamiento y facilitar una lectura integral antes de tomar acciones de seguimiento.")
    add_heading(doc, "Qué puede hacer el docente", 2)
    for text in (
        "Consultar los grupos y alumnos a los que tiene acceso.",
        "Identificar riesgos mediante asistencia, faltas consecutivas, evidencias y pendientes.",
        "Revisar la trayectoria del alumno por periodo y cuatrimestre.",
        "Consultar acuerdos, reportes docentes, sesiones tutoriales y canalizaciones.",
        "Usar la línea de tiempo para reconstruir eventos relevantes en orden cronológico.",
    ): add_bullet(doc, text)
    callout(doc, "IMPORTANTE", "Este módulo es principalmente de consulta. Los datos se originan en asistencia, evaluaciones, tutoría y otros módulos de SIGA. Si falta información, deberá capturarse o corregirse en el módulo correspondiente.", "FFF4E5", AMBER)

    add_heading(doc, "2. Ingreso al expediente", 1)
    add_step(doc, 1, "Inicie sesión", "Acceda a SIGA con su cuenta institucional de docente.")
    add_step(doc, 2, "Abra Mis Tutorados", "En el menú lateral, ubique el grupo Tutoría y seleccione Mis Tutorados.")
    add_step(doc, 3, "Seleccione al estudiante", "Localice al alumno y pulse Ver expediente académico integral.")
    add_step(doc, 4, "Confirme el encabezado", "Verifique nombre, matrícula, carrera, cuatrimestre, grupo, periodo y tutor antes de interpretar los datos.")
    add_heading(doc, "Acceso alternativo", 2)
    doc.add_paragraph("Si su cuenta dispone del permiso de lectura del expediente, también puede abrir directamente la opción Expediente académico y buscar por nombre o matrícula. La búsqueda directa comienza al escribir al menos dos caracteres.")

    doc.add_page_break()
    add_heading(doc, "3. Panorama del grupo", 1)
    doc.add_paragraph("Antes de abrir un expediente individual, el panorama permite priorizar la revisión del grupo. Puede cambiar entre vista de lista y tarjetas, buscar por grupo o carrera, filtrar por cuatrimestre y configuración, y ordenar por grado, carrera o número de alumnos.")
    add_heading(doc, "Indicadores principales", 2)
    rows = [
        ("Asistencia global", "Porcentaje acumulado con los registros disponibles."),
        ("Prom. evidencias", "Promedio interno de evidencias; no representa una calificación oficial."),
        ("En riesgo", "Alumnos con señales que requieren revisión prioritaria."),
        ("Requieren atención", "Alumnos con indicadores preventivos o pendientes."),
        ("Sin información", "Casos sin registros suficientes para una interpretación confiable."),
        ("Cobertura", "Proporción de asistencias capturadas; ayuda a valorar la calidad del panorama."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width, table.columns[1].width = Inches(1.8), Inches(4.7)
    hdr = table.rows[0].cells
    for i, text in enumerate(("Indicador", "Cómo interpretarlo")):
        shade(hdr[i], NAVY); rr = hdr[i].paragraphs[0].add_run(text); rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255)
    set_repeat_header(table.rows[0])
    for a, b in rows:
        cells = table.add_row().cells
        cells[0].width, cells[1].width = Inches(1.8), Inches(4.7)
        cells[0].paragraphs[0].add_run(a).bold = True
        cells[1].paragraphs[0].add_run(b)
    for row in table.rows:
        for cell in row.cells: margins(cell)
    add_heading(doc, "Cómo priorizar", 2)
    for text in (
        "Empiece por los alumnos marcados En riesgo.",
        "Revise después Requieren atención y ordene la lista con los datos de asistencia y pendientes.",
        "En Sin información, compruebe primero la cobertura y la captura de registros; no concluya que el alumno está regular.",
        "Pulse Ver expediente para abrir el detalle individual.",
    ): add_bullet(doc, text)

    add_heading(doc, "4. Consulta del expediente de un alumno", 1)
    doc.add_paragraph("Al abrir un alumno, el encabezado muestra sus datos de identificación y el semáforo general. Debajo aparecen ocho pestañas. Utilice el semáforo como señal de priorización, no como diagnóstico definitivo.")
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width, table.columns[1].width = Inches(1.7), Inches(4.8)
    for i, text in enumerate(("Semáforo", "Lectura recomendada")):
        shade(table.rows[0].cells[i], NAVY); rr=table.rows[0].cells[i].paragraphs[0].add_run(text); rr.bold=True; rr.font.color.rgb=RGBColor(255,255,255)
    for label, desc, fill, color in (
        ("Rojo · Riesgo alto", "Revisar de inmediato las razones, materias, asistencia y pendientes.", "FEE4E2", RED),
        ("Amarillo · Atención", "Analizar tendencias y dar seguimiento preventivo.", "FEF0C7", AMBER),
        ("Verde · Regular", "Mantener monitoreo; no sustituye la revisión profesional.", "DCFCE7", GREEN),
        ("Gris · Información insuficiente", "Validar captura y cobertura antes de emitir conclusiones.", "EAECF0", GRAY),
    ):
        cells=table.add_row().cells; shade(cells[0], fill); r=cells[0].paragraphs[0].add_run(label); r.bold=True; r.font.color.rgb=RGBColor.from_string(color); cells[1].paragraphs[0].add_run(desc)
    for row in table.rows:
        for cell in row.cells: margins(cell)

    doc.add_page_break()
    add_heading(doc, "5. Interpretación de las secciones", 1)
    sections = [
        ("Resumen", "Presenta materias inscritas, asistencia global, promedio de evidencias, materias en riesgo, acuerdos pendientes, reportes abiertos y canalizaciones. Revise las razones del semáforo y las alertas por materia."),
        ("Trayectoria", "Muestra cada inscripción, periodo, cuatrimestre, grupo y resolución de Servicios Escolares. Úsela para comprender continuidad, promoción o antecedentes del alumno."),
        ("Materias", "Compara materia, docente, número de evidencias, promedio, asistencia, faltas y estado. Identifique si el riesgo se concentra en una asignatura."),
        ("Asistencia", "Compara materias y analiza patrones por bloque horario, mapa semanal y días con ausencia parcial. Puede incluir o excluir faltas justificadas del análisis."),
        ("Evaluaciones", "Lista evidencias internas registradas por docentes, con fecha y calificación. Estas cifras son orientativas y no oficiales."),
        ("Acuerdos", "Permite consultar y filtrar acuerdos por materia, estado, tipo y responsable. El docente consulta el seguimiento; la depuración de registros de prueba está reservada a perfiles autorizados."),
        ("Tutoría", "Integra tutor asignado, estado de seguimiento, reportes docentes, sesiones tutoriales y canalizaciones."),
        ("Línea de tiempo", "Ordena eventos de asistencia excepcional, evaluaciones, acuerdos, reportes y tutoría para reconstruir el contexto del caso."),
    ]
    for name, body in sections:
        add_heading(doc, name, 2)
        doc.add_paragraph(body)
    callout(doc, "CRITERIO PROFESIONAL", "Un indicador aislado no explica por sí solo la situación del alumno. Contraste asistencia, evidencias, antecedentes, acuerdos y seguimiento tutorial antes de intervenir.")

    add_heading(doc, "6. Flujo recomendado de seguimiento", 1)
    for num, title, body in (
        (1, "Detectar", "En el panorama del grupo, filtre En riesgo y Requieren atención."),
        (2, "Validar", "Revise cobertura, fechas y cantidad de registros para descartar información incompleta."),
        (3, "Comprender", "Abra Resumen, Materias y Asistencia; después consulte Trayectoria, Acuerdos y Tutoría."),
        (4, "Actuar", "Realice la intervención desde el módulo correspondiente: registro docente, acuerdo, reporte o tutoría, según el procedimiento institucional."),
        (5, "Dar seguimiento", "Consulte nuevamente el expediente en la fecha acordada y verifique cambios en pendientes, asistencia y evidencias."),
        (6, "Escalar", "Cuando el caso lo requiera, canalice por los medios institucionales y evite incluir información sensible innecesaria."),
    ): add_step(doc, num, title, body)

    doc.add_page_break()
    add_heading(doc, "7. Buenas prácticas y confidencialidad", 1)
    for text in (
        "Confirme la identidad del alumno antes de comentar o tomar decisiones.",
        "Consulte únicamente expedientes relacionados con sus funciones.",
        "No comparta capturas, matrículas, calificaciones ni datos tutoriales por canales no autorizados.",
        "Evite dejar el expediente abierto en equipos compartidos; cierre sesión al terminar.",
        "No interprete Sin información como ausencia de riesgo.",
        "Registre hechos verificables y lenguaje respetuoso en los módulos de origen.",
        "No use el promedio de evidencias como calificación oficial.",
        "Si detecta un dato incorrecto, solicite su corrección al área responsable; no compense el error con anotaciones informales.",
    ): add_bullet(doc, text)
    callout(doc, "PRIVACIDAD", "La información académica y tutorial puede contener datos personales. Aplique el principio de mínima divulgación: comparta solo lo necesario, con la persona autorizada y por el canal institucional.", "FEE4E2", RED)

    add_heading(doc, "8. Problemas frecuentes", 1)
    issues = [
        ("No aparece un alumno", "Escriba al menos dos caracteres, pruebe con la matrícula y confirme que el alumno pertenece a un grupo accesible para su cuenta."),
        ("No hay grupos visibles", "Verifique el periodo, las asignaciones y los permisos de su cuenta con el administrador del sistema."),
        ("El grupo aparece “Sin materias”", "La configuración académica aún no tiene materias relacionadas. Solicite la revisión al área responsable."),
        ("El expediente está en gris", "Hay información insuficiente. Revise cobertura, asistencias, clases y evidencias capturadas."),
        ("Los datos parecen desactualizados", "Actualice la página y compruebe si el registro se guardó en el módulo de origen. Si persiste, reporte alumno, matrícula, grupo, sección y fecha del dato."),
        ("No puedo modificar un acuerdo", "El expediente es una vista consolidada. Use el módulo que originó el acuerdo o solicite apoyo al perfil autorizado."),
    ]
    for title, body in issues:
        add_heading(doc, title, 2)
        doc.add_paragraph(body)

    add_heading(doc, "9. Guía rápida", 1)
    callout(doc, "EN 60 SEGUNDOS", "Mis Tutorados → seleccione al alumno → Ver expediente académico integral → confirme identidad → revise semáforo y razones → contraste Materias y Asistencia → consulte Acuerdos y Tutoría → defina seguimiento en el módulo correspondiente.")
    add_heading(doc, "Lista de verificación antes de cerrar", 2)
    for text in (
        "Revisé identidad y periodo del alumno.",
        "Consideré la cobertura y suficiencia de los datos.",
        "Contrasté al menos dos fuentes del expediente.",
        "Identifiqué pendientes, responsable y fecha de revisión.",
        "Protegí la confidencialidad de la información.",
        "Cerré sesión o bloqueé el equipo.",
    ): add_bullet(doc, "☐ " + text)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Fin del manual")
    r.bold = True; r.font.color.rgb = RGBColor.from_string(BLUE)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "Manual docente - Expediente Académico Integral"
    doc.core_properties.subject = "Guía de uso del módulo Expediente Académico Integral de SIGA"
    doc.core_properties.author = "SIGA"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
