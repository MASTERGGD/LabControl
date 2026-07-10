from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/triptico_siga_utecan_inventario.docx")


COLORS = {
    "green": "00855F",
    "green_dark": "00684A",
    "green_light": "EAF7F1",
    "blue": "0F4C81",
    "blue_light": "EAF2FA",
    "gold": "B7791F",
    "gold_light": "FFF7E6",
    "ink": "102033",
    "muted": "52627A",
    "line": "D6E0EA",
    "white": "FFFFFF",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D6E0EA", size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def clear_para(paragraph):
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def add_run(paragraph, text, size=9, bold=False, color=COLORS["ink"]):
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_title(cell, text, subtitle=None, color=COLORS["green_dark"]):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text, size=15, bold=True, color=color)
    if subtitle:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(6)
        add_run(p2, subtitle, size=8.5, color=COLORS["muted"])


def add_kicker(cell, text, fill=COLORS["green_light"], color=COLORS["green_dark"]):
    table = cell.add_table(rows=1, cols=1)
    table.autofit = False
    mini = table.cell(0, 0)
    set_cell_shading(mini, fill)
    set_cell_border(mini, fill, "4")
    set_cell_margins(mini, 60, 100, 60, 100)
    p = mini.paragraphs[0]
    add_run(p, text.upper(), size=7.5, bold=True, color=color)


def add_body(cell, text, size=8.6, bold=False, color=COLORS["ink"], after=4):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    add_run(p, text, size=size, bold=bold, color=color)


def add_bullet(cell, text, color=COLORS["ink"]):
    p = cell.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.02
    add_run(p, "• ", size=8.5, bold=True, color=COLORS["green"])
    add_run(p, text, size=8.3, color=color)


def add_metric(cell, value, label, color=COLORS["green_dark"]):
    t = cell.add_table(rows=1, cols=1)
    t.autofit = False
    c = t.cell(0, 0)
    set_cell_shading(c, COLORS["blue_light"])
    set_cell_border(c, COLORS["line"], "6")
    set_cell_margins(c, 80, 120, 80, 120)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, value, size=18, bold=True, color=color)
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, label, size=7.5, color=COLORS["muted"])


def add_panel_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Inches(3.48), Inches(3.48), Inches(3.48)]
    for i, width in enumerate(widths):
        table.columns[i].width = width
    for cell in table.rows[0].cells:
        cell.width = Inches(3.48)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_margins(cell)
        set_cell_border(cell)
        clear_para(cell.paragraphs[0])
    return table


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(9)
    return doc


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = setup_doc()

    # Exterior: contraportada, portada, contexto.
    exterior = add_panel_table(doc)
    back, cover, context = exterior.rows[0].cells

    set_cell_shading(cover, COLORS["green_light"])
    add_kicker(cover, "Universidad Tecnologica de Candelaria")
    add_title(
        cover,
        "SIGA UTECAN",
        "Sistema Integral de Gestion Academica",
        color=COLORS["green_dark"],
    )
    add_body(cover, "Plataforma institucional para ordenar laboratorios, inventario, mantenimiento, comunicados y procesos academico-administrativos.", size=10.5, bold=True)
    add_body(cover, "Enfoque del demo: Inventario Institucional como base para control patrimonial, trazabilidad y toma de decisiones.", size=9.2, color=COLORS["muted"])
    add_metric(cover, "1", "plataforma integrada")
    add_body(cover, "Preparado para directivos UTECAN", size=8.2, color=COLORS["green_dark"], bold=True)

    add_title(context, "Por que SIGA", "Problemas que ayuda a resolver")
    add_bullet(context, "Informacion dispersa entre archivos, mensajes y registros manuales.")
    add_bullet(context, "Dificultad para saber que activo existe, donde esta y quien lo tiene.")
    add_bullet(context, "Reportes de mantenimiento sin seguimiento claro o sin responsable visible.")
    add_bullet(context, "Procesos que dependen de una sola persona o de evidencia informal.")
    add_kicker(context, "Valor institucional", fill=COLORS["blue_light"], color=COLORS["blue"])
    add_body(context, "SIGA centraliza informacion operativa y deja historial: quien registra, quien valida, quien recibe, quien atiende y cuando se resuelve.", size=8.9)
    add_body(context, "Esto permite pasar de reaccionar a problemas aislados a administrar evidencia para tomar decisiones.", size=8.9)

    add_title(back, "Inversion inicial", "La nube no es gratis: se vuelve sostenible si se planea")
    add_bullet(back, "Base de datos y almacenamiento persistente para documentos, evidencias, QR, respaldos y bitacoras.")
    add_bullet(back, "Servidor de aplicacion, dominio, seguridad, monitoreo y respaldos.")
    add_bullet(back, "Capacitacion inicial y carga/depuracion de catalogos institucionales.")
    add_bullet(back, "Politica de respaldo y crecimiento por volumen de activos, fotos y documentos.")
    add_kicker(back, "Recomendacion")
    add_body(back, "Iniciar con un presupuesto controlado para almacenamiento y respaldos, medir uso real durante el piloto y ajustar el plan antes de escalar a toda la universidad.", size=8.8)
    add_body(back, "Siguiente paso: validar flujo de inventario con Finanzas, Responsables de Departamento, Sistemas e Infraestructura.", size=8.5, bold=True, color=COLORS["green_dark"])

    doc.add_page_break()

    # Interior: inventario, flujo, beneficios/retos.
    interior = add_panel_table(doc)
    inventory, flow, benefits = interior.rows[0].cells

    set_cell_shading(inventory, COLORS["gold_light"])
    add_kicker(inventory, "Modulo principal del demo", fill=COLORS["white"], color=COLORS["gold"])
    add_title(inventory, "Inventario Institucional", "Control patrimonial con trazabilidad")
    add_body(inventory, "Permite registrar, revisar, validar y consultar activos por departamento, laboratorio, categoria, responsable y estado.", size=9.2, bold=True)
    add_bullet(inventory, "Codigo SIGA para identificacion interna y numero patrimonial cuando exista.")
    add_bullet(inventory, "Flujo de borrador, revision, validacion, observacion y no autorizacion.")
    add_bullet(inventory, "Transferencias entre departamentos con aceptacion del receptor.")
    add_bullet(inventory, "Expediente del bien: historial, movimientos, resguardante, ubicacion y QR.")
    add_bullet(inventory, "Panel ejecutivo: categorias, responsables, pendientes y alertas.")

    add_title(flow, "Flujo sugerido", "De captura a inventario oficial")
    add_body(flow, "1. Captura", size=8.8, bold=True, color=COLORS["green_dark"])
    add_body(flow, "Un capturista o responsable registra activos por Excel o manualmente.", size=8.2)
    add_body(flow, "2. Revision", size=8.8, bold=True, color=COLORS["green_dark"])
    add_body(flow, "El responsable valida datos: categoria, ubicacion, resguardante, serie, estado y evidencia.", size=8.2)
    add_body(flow, "3. Validacion", size=8.8, bold=True, color=COLORS["green_dark"])
    add_body(flow, "El activo entra al inventario operativo y habilita etiqueta QR, movimientos, resguardos y mantenimiento.", size=8.2)
    add_body(flow, "4. Movimiento", size=8.8, bold=True, color=COLORS["green_dark"])
    add_body(flow, "Transferencias y cambios quedan registrados; el departamento receptor acepta o rechaza.", size=8.2)
    add_kicker(flow, "Resultado")
    add_body(flow, "Cada bien conserva su historia aun cuando cambie de departamento, ubicacion o resguardante.", size=8.7)

    add_title(benefits, "Ventajas y retos", "Lo que direccion debe considerar")
    add_body(benefits, "Ventajas", size=9, bold=True, color=COLORS["green_dark"])
    add_bullet(benefits, "Mayor control del patrimonio universitario.")
    add_bullet(benefits, "Responsabilidades claras por departamento.")
    add_bullet(benefits, "Menos perdida de informacion y mejor evidencia para auditorias.")
    add_bullet(benefits, "Base para mantenimiento, bajas, prestamos y reemplazos.")
    add_body(benefits, "Retos de implementacion", size=9, bold=True, color=COLORS["gold"])
    add_bullet(benefits, "Depurar catalogos: departamentos, usuarios, laboratorios y ubicaciones.")
    add_bullet(benefits, "Definir responsables y permisos reales.")
    add_bullet(benefits, "Estandarizar codigos, resguardos, etiquetas y evidencia.")
    add_bullet(benefits, "Asignar almacenamiento y respaldo desde el inicio.")
    add_kicker(benefits, "Mensaje clave", fill=COLORS["green_light"], color=COLORS["green_dark"])
    add_body(benefits, "SIGA no sustituye las decisiones institucionales: las ordena, documenta y vuelve medibles.", size=9, bold=True)

    doc.save(OUT)


if __name__ == "__main__":
    build()
